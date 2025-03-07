import re
import spacy
import phonenumbers
import docx
import PyPDF2
import pdfplumber
import os
from natasha import (
    Segmenter,
    MorphVocab,
    NewsEmbedding,
    NewsMorphTagger,
    NewsSyntaxParser,
    NewsNERTagger,
    NamesExtractor,
    Doc
)
from transformers import pipeline
import torch
import logging
from typing import List, Dict, Optional

logger = logging.getLogger(__name__)

class AdvancedResumeParser:
    def __init__(self):
        logger.info("Initializing resume parser...")
        
        # Initialize Natasha
        self.segmenter = Segmenter()
        self.morph_vocab = MorphVocab()
        self.emb = NewsEmbedding()
        self.morph_tagger = NewsMorphTagger(self.emb)
        self.syntax_parser = NewsSyntaxParser(self.emb)
        self.ner_tagger = NewsNERTagger(self.emb)
        self.names_extractor = NamesExtractor(self.morph_vocab)
        
        # Initialize Spacy
        self.nlp = spacy.load("ru_core_news_sm")
        
        # Initialize text classification model
        try:
            self.classifier = pipeline(
                "zero-shot-classification",
                model="facebook/bart-large-mnli",
                device=0 if torch.cuda.is_available() else -1
            )
            logger.info("Classification model successfully loaded")
        except Exception as e:
            logger.error(f"Error loading classification model: {e}")
            self.classifier = None
        
        # Regular expressions
        self.email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        
        # Dictionary of specializations and keywords
        self.specializations = {
            'QA': ['qa', 'quality assurance', 'qa engineer', 'qa automation', 'qa tester', 'qa testing', 'qa control', 'qa assurance', 'qa assurance engineer', 'qa assurance tester', 'qa assurance automation', 'qa assurance testing', 'тестировщик', 'тестирование', 'quality control'],
            'Python': ['python', 'django', 'flask', 'fastapi', 'pytest'],
            'Java': ['java', 'spring', 'hibernate', 'maven'],
            'Frontend': ['javascript', 'react', 'vue', 'angular', 'frontend', 'фронтенд'],
            'Backend': ['backend', 'бэкенд', 'node.js', 'php', 'ruby'],
            'DevOps': ['devops', 'docker', 'kubernetes', 'ci/cd', 'jenkins'],
            'Android': ['android', 'kotlin', 'mobile'],
            'iOS': ['ios', 'swift', 'objective-c'],
            'Data Science': ['data science', 'data scientist', 'machine learning', 'ml engineer', 'nlp', 'sklearn', 'scikit-learn', 'pytorch'],
            'Business Analyst': ['business analyst', 'бизнес-аналитик', 'аналитик', 'анализ требований'],
            'System Analyst': ['system analyst', 'системный аналитик', 'системный архитектор'],
            'Analyst': ['analyst', 'аналитик', 'data analyst', 'data analysis'],
            'Product Manager': ['product manager', 'продакт-менеджер', 'продукт-менеджер', 'управление продуктом', 'product owner'],
            'Project Manager': ['project manager', 'проектный менеджер', 'управление проектами', 'pmbok', 'scrum', 'agile', 'канбан'],
            'Designer': ['designer', 'дизайнер', 'figma', 'sketch', 'photoshop', 'ui/ux', 'web design', 'графический дизайн'],
        }

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extracts text from PDF using multiple methods."""
        text = ""
        # Try pdfplumber first for better text extraction
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            logger.warning(f"Error reading PDF with pdfplumber: {e}")
            # Try PyPDF2 as fallback
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() or ""
            except Exception as e:
                logger.error(f"Error reading PDF with PyPDF2: {e}")
        
        return text

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extracts text from DOCX file."""
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Add text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
        return text

    def extract_text(self, file_path: str) -> str:
        """Extracts text from file based on its type."""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            logger.error(f"Unsupported file extension: {file_extension}")
            return ""

    def extract_name(self, text: str) -> Optional[str]:
        """Extracts name from text using Natasha and Spacy."""

        extracted_name = None
        # for hh.ru resumes
        if "Желаемая должность и зарплата" in text:
            extracted_name = text.split("\n")[0]
            # return text.split("\n")[0]

        if extracted_name is None:
            # Try Spacy first
            try:
                doc_spacy = self.nlp(' . '.join(text[:100].split('\n')))  # Analyze first 100 characters
                for ent in doc_spacy.ents:
                    if ent.label_ == "PER":
                        extracted_name = ent.text
                        break
            except Exception as e:
                logger.warning(f"Error processing name with Spacy: {e}")
        

        if extracted_name is None:
            try:
                # Split text into lines and find first non-empty line
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                if lines:
                    # Take first two words from first non-empty line
                    words = lines[0].split()
                    if len(words) >= 2:
                        extracted_name = " ".join(words[:2])
            except Exception as e:
                logger.warning(f"Error extracting name from text: {e}")
        
        if extracted_name is not None:
            extracted_name = extracted_name.lower()
            make_title = lambda x: x.group(0).title()
            extracted_name = re.sub(r"[а-яА-Яa-zA-Z]+", make_title, extracted_name)
            
        return extracted_name

    def extract_phone_numbers(self, text: str) -> List[str]:
        """Extracts phone numbers from text."""
        phones = []
        # Look for phone numbers in text
        for match in phonenumbers.PhoneNumberMatcher(text, "RU"):
            phone = phonenumbers.format_number(match.number, phonenumbers.PhoneNumberFormat.INTERNATIONAL)
            phones.append(phone)
        return phones

    def extract_email(self, text: str) -> Optional[str]:
        """Extracts email from text."""
        emails = re.findall(self.email_pattern, text)
        return emails[0] if emails else None

    def extract_specialization_ai(self, text: str) -> List[str]:
        """Uses model to determine specialization."""
        if not self.classifier:
            return []
        
        try:
            # Prepare list of possible specializations
            candidate_labels = list(self.specializations.keys())
            
            # Classify text
            result = self.classifier(
                text[:500],  # Take first 500 characters
                candidate_labels,
                multi_label=True
            )
            
            # Filter results with high confidence (more than 0.85)
            logger.info(f"Extracted specializations AI: {list(zip(result['labels'], result['scores']))}")
            return [label for label, score in zip(result['labels'], result['scores']) if score > 0.85]
        except Exception as e:
            logger.error(f"Error using classifier: {e}")
            return []

    def extract_specialization(self, text: str) -> List[str]:
        """Extracts specializations by analyzing the beginning of the resume."""
        # Take only first 500 characters for analysis
        text_beginning = text[:500].lower()
        found_specs = []
        
        # Basic search by keywords
        for spec, keywords in self.specializations.items():
            for keyword in keywords:
                if keyword.lower() in text_beginning:
                    found_specs.append(spec)
                    break
        
        # If no specializations found, try AI classification
        if not found_specs:
            try:
                # Prepare list of possible specializations
                candidate_labels = list(self.specializations.keys())
                
                # Classify text
                result = self.classifier(
                    text_beginning,
                    candidate_labels,
                    multi_label=True
                )
                
                # Filter results with high confidence (more than 0.85)
                logger.info(f"Extracted specializations: {list(zip(result['labels'], result['scores']))}")
                ai_specs = [label for label, score in zip(result['labels'], result['scores']) if score > 0.85]
                found_specs.extend(ai_specs)
            except Exception as e:
                logger.error(f"Error using classifier: {e}")
        
        # Remove duplicates and return unique specializations
        return list(set(found_specs))

    def parse_resume(self, file_path: str) -> Optional[Dict]:
        """Main method for parsing resume."""
        logger.info(f"Starting to parse file: {file_path}")
        
        text = self.extract_text(file_path)
        if not text:
            logger.warning("Failed to extract text from file")
            return None

        # Extract all data
        name = self.extract_name(text)
        phones = self.extract_phone_numbers(text)
        email = self.extract_email(text)
        
        # Extract specializations from the beginning of the resume
        specializations = self.extract_specialization(text)
        
        logger.info(f"Found: name={name}, phones={phones}, email={email}, specializations={specializations}")

        return {
            'name': name,
            'email': email,
            'phone': phones[0] if phones else None,
            'specializations': specializations,
            'raw_text': text  # Save original text for further analysis if needed
        } 